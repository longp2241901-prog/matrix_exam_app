import streamlit as st
from groq import Groq
import os
import re
import random
from docx import Document
from docx.shared import RGBColor
import PyPDF2
from io import BytesIO
from docx import Document as DocReader
import copy
import zipfile
#=====================
# ğŸ”“ Giáº£i nÃ©n data.zip náº¿u chÆ°a cÃ³ thÆ° má»¥c data
if not os.path.exists("data"):
    if os.path.exists("data.zip"):
        with zipfile.ZipFile("data.zip", 'r') as zip_ref:
            zip_ref.extractall(".")
        print("âœ… ÄÃ£ giáº£i nÃ©n data.zip")
    else:
        print("âš ï¸ KhÃ´ng tÃ¬m tháº¥y data.zip")
# =========================
# ğŸ§¹ HÃ m lÃ m sáº¡ch ná»™i dung trÆ°á»›c khi Tex hÃ³a
# =========================
def clean_text_for_tex(text: str) -> str:
    """Bá» 'CÃ¢u 1.', 'A.', 'B.'... vÃ  lÃ m gá»n vÄƒn báº£n"""
    # Bá» CÃ¢u 1., CÃ¢u 2.
    text = re.sub(r"C[Ã¢a]u\s*\d+\s*[.:]", "", text, flags=re.IGNORECASE)
    # Bá» A. B. C. D. (tráº¯c nghiá»‡m)
    text = re.sub(r"\b[ABCDÄ]\s*\.", "", text)
    # Bá» a) b) c) d) (Ä‘Ãºng/sai)
    text = re.sub(r"\b[a-d]\)", "", text)
    # LÃ m gá»n khoáº£ng tráº¯ng
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


# =========================
# âš™ï¸ Cáº¥u hÃ¬nh trang
# =========================
st.set_page_config(layout="wide")
#st.title("ğŸ“ Sinh Ä‘á» kiá»ƒm tra tá»« ma tráº­n (chuáº©n ex_test)")
# =========================   
# ğŸ§® ThÃ´ng tin á»©ng dá»¥ng & TÃ¡c giáº£ (hiá»ƒn thá»‹ Ä‘áº§u trang)
# =========================
st.markdown(
    """
    <div style='text-align: center; line-height: 1.6; margin-bottom: 20px;'>
        <img src="https://cdn-icons-png.flaticon.com/512/3523/3523063.png" width="55" style="margin-bottom: 5px;" />
        <h1 style="margin-bottom: 0;">SinhÄá»+</h1>
        <p style="color: gray; font-size: 16px; margin-top: 4px;">
            á»¨ng dá»¥ng sinh Ä‘á» kiá»ƒm tra tá»± Ä‘á»™ng â€” <b>Pháº¡m Tiáº¿n Long & TrÆ°Æ¡ng Thá»‹ Huá»³nh Trang</b> (2025)
        </p>
        <p style="font-size: 15px; color: #555;">
            ğŸ“ LiÃªn há»‡ há»— trá»£: <a href="tel:0396595129" style="text-decoration: none; color: #3366cc;">0396595129</a><br>
            âœ‰ï¸ Email: <a href="mailto:longp2241901@gmail.com" style="text-decoration: none; color: #3366cc;">longp2241901@gmail.com</a>
        </p>
    </div>
    """,
    unsafe_allow_html=True
)



# =========================
# ğŸ”‘ Nháº­p API Key
# =========================
# =========================
# =========================
# ğŸ”‘ Nháº­p Groq API Key cÃ¡ nhÃ¢n
# =========================
st.markdown("### ğŸ” Nháº­p key Groq API cÃ¡ nhÃ¢n")

# Ã” nháº­p API key
user_api_key = st.text_input(
    "Nháº­p Groq API Key cá»§a báº¡n (báº¯t Ä‘áº§u báº±ng 'gsk_...')",
    type="password",
    help="Báº¡n cáº§n cÃ³ Groq API Key riÃªng Ä‘á»ƒ sá»­ dá»¥ng. Láº¥y táº¡i https://console.groq.com/keys",
)

# HÆ°á»›ng dáº«n thÃªm
st.info(
    """
    ğŸ’¡ **CÃ¡ch láº¥y Groq API Key:**
    1. Truy cáº­p [https://console.groq.com/keys](https://console.groq.com/keys)
    2. ÄÄƒng nháº­p (hoáº·c táº¡o tÃ i khoáº£n miá»…n phÃ­)
    3. Chá»n **Create API Key**
    4. Sao chÃ©p key (dáº¡ng `gsk_...`) vÃ  dÃ¡n vÃ o Ã´ trÃªn.
    
    âš ï¸ **LÆ°u Ã½ giá»›i háº¡n sá»­ dá»¥ng:**
    - Má»—i API key cÃ³ giá»›i háº¡n ~100.000 token má»—i ngÃ y (Ä‘áº¿m cáº£ input + output).  
    - Náº¿u vÆ°á»£t giá»›i háº¡n, báº¡n sáº½ tháº¥y lá»—i `Rate limit reached`.  
    - Sau khoáº£ng **30â€“60 phÃºt**, Groq sáº½ tá»± Ä‘á»™ng reset quota Ä‘á»ƒ báº¡n tiáº¿p tá»¥c sá»­ dá»¥ng.
    """,
    icon="â„¹ï¸"
)

# LÆ°u key vÃ o session
if user_api_key:
    st.session_state["api_key"] = user_api_key.strip()
    st.success("âœ… API Key Ä‘Ã£ Ä‘Æ°á»£c lÆ°u. Báº¡n cÃ³ thá»ƒ báº¯t Ä‘áº§u sá»­ dá»¥ng á»©ng dá»¥ng.")
else:
    st.warning("ğŸ”‘ HÃ£y nháº­p API Key Ä‘á»ƒ tiáº¿p tá»¥c.")

# Náº¿u chÆ°a cÃ³ key thÃ¬ dá»«ng app
if "api_key" not in st.session_state:
    st.stop()

# GÃ¡n biáº¿n dÃ¹ng chung cho toÃ n app
api_key = st.session_state["api_key"]



# =========================
# ğŸ§  HÃ m tiá»‡n Ã­ch
# =========================
def get_sample_file(mon, lop, topic, dang_cauhoi, muc_do, dang):
    base_dir = "data"
    folder = os.path.join(base_dir, mon, lop, topic, dang_cauhoi, muc_do)
    filename = f"{dang}.txt"
    return os.path.join(folder, filename)

def split_ex_blocks(latex_text):
    """TÃ¡ch tá»«ng cÃ¢u há»i \\begin{ex} ... \\end{ex}"""
    return re.findall(r"\\begin{ex}.*?\\end{ex}", latex_text, re.S)

# =========================
# ğŸ’¾ Xuáº¥t LaTeX
# =========================
def export_latex_ex(all_questions, filename="output.tex"):
    latex_content = (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[utf8]{vietnam}\n"
        "\\usepackage{ex_test}\n"
        "\\begin{document}\n"
        "\\section*{Äá» kiá»ƒm tra}\n"
    )
    latex_content += "\n\n".join(all_questions)
    latex_content += "\n\\end{document}"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(latex_content)
    return filename

# =========================
# ğŸ’¾ Xuáº¥t Word
# =========================
def export_word_ex(all_questions, filename="output.docx"):
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Äá» kiá»ƒm tra", 0)
    questions = []
    for q in all_questions:
        questions.extend(split_ex_blocks(q))

    # --- TÃ¡ch cÃ¢u theo loáº¡i ---
    mc_questions = []      # \choice
    tf_questions = []      # \choiceTF
    short_questions = []   # \shortans

    for q in questions:
        if "\\choice" in q and not "\\choiceTF" in q:
            mc_questions.append(q)
        elif "\\choiceTF" in q:
            tf_questions.append(q)
        elif "\\shortans" in q:
            short_questions.append(q)

    section_map = [
        ("PHáº¦N I â€“ TRáº®C NGHIá»†M 4 Lá»°A CHá»ŒN", mc_questions),
        ("PHáº¦N II â€“ TRáº®C NGHIá»†M ÄÃšNG SAI", tf_questions),
        ("PHáº¦N III â€“ TRáº¢ Lá»œI NGáº®N", short_questions)
    ]

    ques_counter = 1
    for title, qlist in section_map:
        if not qlist:
            continue

        # TiÃªu Ä‘á» section
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(14)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")  # thÃªm 1 dÃ²ng trá»‘ng

        for q in qlist:
            # ===== Ná»™i dung cÃ¢u há»i =====
            noi_dung_match = re.search(
                r"\\begin\{ex\}([\s\S]*?)(?=\\choice|\\choiceTF|\\shortans|\\loigiai|\\end\{ex\})",
                q, re.MULTILINE,
            )
            noi_dung = noi_dung_match.group(1).strip() if noi_dung_match else q
            noi_dung = noi_dung.replace("\\\\", "\n").replace("\r", "")

            p = doc.add_paragraph()
            run_q = p.add_run(f"CÃ¢u {ques_counter}. ")
            run_q.bold = True
            p.add_run(noi_dung)

            dap_an = None

            # ===== Tráº¯c nghiá»‡m nhiá»u lá»±a chá»n =====
            if "\\choice" in q and not "\\choiceTF" in q:
                lc_block = re.search(r"\\choice(.*?)(?=\\loigiai|\\end{ex})", q, re.S)
                if lc_block:
                    lines = lc_block.group(1).splitlines()
                    options = []
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        is_true = "\\True" in line
                        line = line.replace("\\True", "").strip("{} ")
                        options.append((line, is_true))
                    for j, (opt, is_true) in enumerate(options):
                        label = chr(65 + j) + "."
                        p_opt = doc.add_paragraph()
                        run = p_opt.add_run(f"{label} {opt}")
                        if is_true:
                            run.bold = True
                            run.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            dap_an = chr(65 + j)

            # ===== ÄÃºng / Sai =====
            elif "\\choiceTF" in q:
                tf_block = re.search(r"\\choiceTF(.*?)(?=\\loigiai|\\end{ex})", q, re.S)
                if tf_block:
                    lines = tf_block.group(1).splitlines()
                    tf_ans = ""
                    idx_tf = 0
                    for line in lines:
                        line = line.strip("{} \t")
                        if not line:
                            continue
                        is_true = "\\True" in line
                        clean_line = line.replace("\\True", "").strip()
                        label = f"{chr(97 + idx_tf)})"
                        p_opt = doc.add_paragraph()
                        run = p_opt.add_run(f"{label} {clean_line}")
                        if is_true:
                            run.bold = True
                            run.underline = True
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        tf_ans += "Ä" if is_true else "S"
                        idx_tf += 1
                    dap_an = tf_ans

            # ===== Tráº£ lá»i ngáº¯n =====
            elif "\\shortans" in q:
                sa_block = re.search(r"\\shortans\{(.*?)\}", q)
                if sa_block:
                    doc.add_paragraph("Tráº£ lá»i ngáº¯n: ............")
                    dap_an = sa_block.group(1).strip()

            # ===== Lá»i giáº£i =====
            loi_giai_match = re.search(r"\\loigiai\{([\s\S]*?)(?=\\end\{ex\})", q)
            if loi_giai_match:
                loi_giai = loi_giai_match.group(1).strip()
                loi_giai = loi_giai.replace("\\\\", "\n").strip()
                if loi_giai.endswith("}"):
                    loi_giai = loi_giai[:-1].rstrip()

                p_lg = doc.add_paragraph()
                run_lg = p_lg.add_run("Lá»i giáº£i: ")
                run_lg.bold = True
                if dap_an:
                    p_lg.add_run(f"ÄÃ¡p Ã¡n: {dap_an}. {loi_giai}")
                else:
                    p_lg.add_run(loi_giai)
            else:
                if dap_an:
                    p_lg = doc.add_paragraph()
                    run_lg = p_lg.add_run("Lá»i giáº£i: ")
                    run_lg.bold = True
                    p_lg.add_run(f"ÄÃ¡p Ã¡n: {dap_an}.")

            ques_counter += 1

    doc.save(filename)
    return filename




# =========================
# âš™ï¸ Cháº¿ Ä‘á»™ nháº­p dá»¯ liá»‡u
# =========================
mode = st.radio(
    "Chá»n cháº¿ Ä‘á»™ lÃ m viá»‡c:",
    [
        "ğŸ“‚ DÃ¹ng dá»¯ liá»‡u cÃ³ sáºµn trong thÆ° má»¥c data",
        "âœï¸ Nháº­p cÃ¢u há»i máº«u thá»§ cÃ´ng",
        "ğŸ“¤ KÃ©o tháº£ file PDF"
    ],
    horizontal=True
)

# =========================
# ğŸ“‚ Giao diá»‡n cÅ© - dÃ¹ng data
# =========================
# =========================
# ğŸ“‚ Giao diá»‡n cÅ© - dÃ¹ng data (má»Ÿ rá»™ng thÃªm mÃ´n)
# =========================
if mode.startswith("ğŸ“‚"):
    def list_subfolders(path):
        return [f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))] if os.path.exists(path) else []
    def list_txt_files(path):
        return [f[:-4] for f in os.listdir(path) if f.endswith(".txt")] if os.path.exists(path) else []
    BASE_DIR = "data"
    st.markdown("## ğŸ§© Ma tráº­n chá»n cÃ¢u há»i")
    ALL_MON = sorted(list_subfolders(BASE_DIR)) if os.path.exists(BASE_DIR) else []
    
    if "configs" not in st.session_state:
        st.session_state.configs = [{"mon": "", "lop": "", "topic": "", "dang_cauhoi": "", "muc_do": "", "dang": "", "count": 1}]
    
    if st.button("â• ThÃªm cáº¥u hÃ¬nh"):
        st.session_state.configs.append({"mon": "", "lop": "", "topic": "", "dang_cauhoi": "", "muc_do": "", "dang": "", "count": 1})
        st.rerun()

    for idx, cfg in enumerate(list(st.session_state.configs)):
        cols = st.columns([1.2,1.2,1.6,1.4,1.4,1.6,0.9,0.8])
        
        # ğŸ”¹ MÃ´n
        with cols[0]:
            mon_folders = list_subfolders(BASE_DIR)
            cfg["mon"] = st.selectbox("MÃ´n", mon_folders, key=f"mon_{idx}") if mon_folders else ""
        
        # ğŸ”¹ Lá»›p
        with cols[1]:
            lops = list_subfolders(os.path.join(BASE_DIR, cfg["mon"])) if cfg["mon"] else []
            cfg["lop"] = st.selectbox("Lá»›p", lops, key=f"lop_{idx}") if lops else ""
        
        # ğŸ”¹ Chá»§ Ä‘á»
        with cols[2]:
            topics = list_subfolders(os.path.join(BASE_DIR, cfg["mon"], cfg["lop"])) if cfg["lop"] else []
            cfg["topic"] = st.selectbox("Chá»§ Ä‘á»", topics, key=f"topic_{idx}") if topics else ""
        
        # ğŸ”¹ Loáº¡i cÃ¢u há»i
        with cols[3]:
            dang_cauhoi = list_subfolders(os.path.join(BASE_DIR, cfg["mon"], cfg["lop"], cfg["topic"])) if cfg["topic"] else []
            cfg["dang_cauhoi"] = st.selectbox("Loáº¡i", dang_cauhoi, key=f"dang_{idx}") if dang_cauhoi else ""
        
        # ğŸ”¹ Má»©c Ä‘á»™
        with cols[4]:
            mucdos = list_subfolders(os.path.join(BASE_DIR, cfg["mon"], cfg["lop"], cfg["topic"], cfg["dang_cauhoi"])) if cfg["dang_cauhoi"] else []
            cfg["muc_do"] = st.selectbox("Má»©c Ä‘á»™", mucdos, key=f"mucdo_{idx}") if mucdos else ""
        
        # ğŸ”¹ Dáº¡ng
        with cols[5]:
            dang_files = list_txt_files(os.path.join(BASE_DIR, cfg["mon"], cfg["lop"], cfg["topic"], cfg["dang_cauhoi"], cfg["muc_do"])) if cfg["muc_do"] else []
            cfg["dang"] = st.selectbox("Dáº¡ng", dang_files, key=f"file_{idx}") if dang_files else ""
        
        # ğŸ”¹ Sá»‘ lÆ°á»£ng
        with cols[6]:
            cfg["count"] = st.number_input("Sá»‘ lÆ°á»£ng", 1, 50, cfg.get("count", 1), key=f"count_{idx}")
        
        # ğŸ”¹ XÃ³a cáº¥u hÃ¬nh
        with cols[7]:
            if st.button("âŒ", key=f"remove_{idx}"):
                st.session_state.configs.pop(idx)
                st.rerun()


# =========================
# âœï¸ Giao diá»‡n nháº­p tay
# =========================
elif mode.startswith("âœï¸"):
    st.markdown("## âœï¸ Nháº­p ná»™i dung cÃ¢u há»i máº«u (theo chuáº©n ex_test)")
    user_input = st.text_area(
        "Nháº­p ná»™i dung LaTeX cá»§a cÃ¢u há»i (\\begin{ex} ... \\end{ex}):",
        height=300,
        placeholder="""VÃ­ dá»¥:
Dáº¡ng 4 lá»±a chá»n: \\begin{ex} ... \\choice{A}{\\True B}{C}{D} \\loigiai{Giáº£i thÃ­ch...} \\end{ex}
Dáº¡ng Ä‘Ãºng sai: \\begin{ex} ... \\choiceTF{a}{\\True b}{c}{\\True d} \\loigiai{Giáº£i thÃ­ch...} \\end{ex}
Dáº¡ng tráº£ lá»i ngáº¯n: \\begin{ex} ... \\shortans[oly]{Ä‘Ã¡p sá»‘}\\end{ex}
Hoáº·c tá»± nháº­p vá»›i ná»™i dung mong muá»‘n, vÃ­ dá»¥: Táº¡o cho tÃ´i bÃ i táº­p thá»‘ng kÃª tÃ­nh tá»© phÃ¢n vá»‹ (ToÃ¡n lá»›p 10)
"""
    )
    so_luong_tu_nhap = st.number_input("Sá»‘ lÆ°á»£ng cÃ¢u muá»‘n sinh thÃªm:", 1, 50, 5)

# =========================
# ğŸ“¤ KÃ©o tháº£ Word / PDF
# =========================
else:
    st.markdown("## ğŸ“¤ KÃ©o tháº£ file PDF Ä‘á»ƒ Ä‘á»c ná»™i dung")
    st.info(
        """
        ğŸ’¡ **HÆ°á»›ng dáº«n sá»­ dá»¥ng:**
        - á»¨ng dá»¥ng chá»‰ há»— trá»£ **file PDF**.
        - Náº¿u báº¡n cÃ³ file **Word (.docx)** chá»©a Ä‘á» gá»‘c, vui lÃ²ng **chuyá»ƒn sang PDF** trÆ°á»›c khi táº£i lÃªn.
        - CÃ¡ch Ä‘Æ¡n giáº£n nháº¥t: Má»Ÿ Word â†’ Chá»n **File â†’ Save As â†’ PDF**.
        - Sau khi táº£i lÃªn PDF, há»‡ thá»‘ng sáº½ tá»± Ä‘á»™ng Ä‘á»c, lÃ m sáº¡ch vÃ  Tex hÃ³a ná»™i dung.
        - Äá»‘i vá»›i file word sau khi xuáº¥t ra, Ä‘á»ƒ chuyá»ƒn cÃ´ng thá»©c ToÃ¡n sang Mathtype thÃ¬ chá»n cÃ´ng thá»©c sau Ä‘Ã³ chá»n Mathtype->Toogle TeX.
        âš ï¸ Má»—i láº§n xá»­ lÃ½, á»©ng dá»¥ng chá»‰ Ä‘á»c **tá»‘i Ä‘a 2 trang Ä‘áº§u tiÃªn cá»§a PDF** Ä‘á»ƒ Ä‘áº£m báº£o tá»‘c Ä‘á»™ vÃ  Ä‘á»™ chÃ­nh xÃ¡c.
        âš ï¸ Báº¡n cÃ³ thá»ƒ dÃ¹ng **khoáº£ng 10â€“12 láº§n/ngÃ y** trÆ°á»›c khi Ä‘áº¡t giá»›i háº¡n token. Khi Ä‘áº¡t giá»›i háº¡n token hÃ£y **chá» 30â€“60 phÃºt** Ä‘á»ƒ tiáº¿p tá»¥c.
        """,
        icon="â„¹ï¸"
    )
    uploaded_file = st.file_uploader("ğŸ“„ KÃ©o tháº£ hoáº·c chá»n file PDF táº¡i Ä‘Ã¢y", type=["pdf"])
    extracted_text = ""
#==========
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()
        extracted_text = ""

        if file_type == "docx":
            doc = DocReader(uploaded_file)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        elif file_type == "pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            total_pages = len(pdf_reader.pages)

            # âœ… Cho phÃ©p ngÆ°á»i dÃ¹ng chá»n trang (vÃ­ dá»¥: "1,2" hoáº·c "5-6")
            page_input = st.text_input(
                f"Nháº­p sá»‘ trang cáº§n Tex hÃ³a (1â€“{total_pages}, tá»‘i Ä‘a 2 trang):",
                value="1,2"
            )

            # ğŸ”¢ HÃ m láº¥y danh sÃ¡ch trang tá»« chuá»—i nháº­p
            def parse_page_input(text):
                pages = set()
                for part in text.split(","):
                    part = part.strip()
                    if "-" in part:
                        start, end = part.split("-")
                        pages.update(range(int(start), int(end) + 1))
                    elif part.isdigit():
                        pages.add(int(part))
                # Giá»›i háº¡n tá»‘i Ä‘a 2 trang
                return sorted(list(pages))[:2]

            selected_pages = parse_page_input(page_input)
            selected_pages = [p for p in selected_pages if 1 <= p <= total_pages]

            if not selected_pages:
                st.warning("âš ï¸ Vui lÃ²ng nháº­p sá»‘ trang há»£p lá»‡ (tá»‘i Ä‘a 2 trang).")
            else:
                st.info(f"ğŸ“„ Äang Ä‘á»c cÃ¡c trang: {', '.join(map(str, selected_pages))}")
                extracted_text = ""
                for p in selected_pages:
                    page = pdf_reader.pages[p - 1]
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"

        # ğŸ”¹ LÃ m sáº¡ch ná»™i dung
        extracted_text = clean_text_for_tex(extracted_text)

        st.text_area("ğŸ“œ Ná»™i dung Ä‘á»c Ä‘Æ°á»£c:", extracted_text, height=300)

    
#=======
        action = st.radio("Chá»n hÃ nh Ä‘á»™ng:", ["ğŸ§  Tex hÃ³a ná»™i dung", "ğŸš€ Sinh Ä‘á» tÆ°Æ¡ng tá»±"], horizontal=True)

        if st.button("âš™ï¸ Thá»±c hiá»‡n"):
            client = Groq(api_key=api_key)
            if action.startswith("ğŸ§ "):
            #====
                prompt = f"""
HÃ£y chuyá»ƒn vÄƒn báº£n sau Ä‘Ã¢y thÃ nh Ä‘á»‹nh dáº¡ng LaTeX theo chuáº©n ex_test.

YÃªu cáº§u:
- KhÃ´ng thÃªm 'CÃ¢u 1.' hoáº·c 'CÃ¢u 2.'.
- Náº¿u cÃ³ cÃ¡c lá»±a chá»n tráº¯c nghiá»‡m (A., B., C., D.), hÃ£y chuyá»ƒn thÃ nh:
  \\choice
  {{Ä‘Ã¡p Ã¡n 1}}
  {{Ä‘Ã¡p Ã¡n 2}}
  {{Ä‘Ã¡p Ã¡n 3}}
  {{Ä‘Ã¡p Ã¡n 4}}
   (má»—i Ä‘Ã¡p Ã¡n trÃªn 1 dÃ²ng riÃªng)
- Náº¿u lÃ  bÃ i Ä‘Ãºng/sai, dÃ¹ng:
  \\choiceTF
  {{má»‡nh Ä‘á» 1}}
  {{má»‡nh Ä‘á» 2}}
  {{má»‡nh Ä‘á» 3}}
  {{má»‡nh Ä‘á» 4}}
- Má»—i bÃ i Ä‘áº·t trong \\begin{{ex}} ... \\end{{ex}}, cÃ³ \\loigiai{{...}} á»Ÿ cuá»‘i.
VÄƒn báº£n cáº§n xá»­ lÃ½:
{extracted_text}
âš ï¸ Chá»‰ tráº£ vá» LaTeX thuáº§n, khÃ´ng thÃªm lá»i giáº£i thÃ­ch.
"""
            else:
                prompt = f"""
DÆ°á»›i Ä‘Ã¢y lÃ  ná»™i dung vÄƒn báº£n ngÆ°á»i dÃ¹ng cung cáº¥p:
{extracted_text}

HÃ£y sinh ra **sá»‘ lÆ°á»£ng cÃ¢u há»i tÆ°Æ¡ng Ä‘Æ°Æ¡ng** vá»›i sá»‘ lÆ°á»£ng cÃ¢u há»i trong vÄƒn báº£n trÃªn (giá»‘ng phong cÃ¡ch, chá»§ Ä‘á», Ä‘á»™ khÃ³, Ä‘á»™ dÃ i) 
vÃ  TRáº¢ Vá»€ Ä‘Ãºng chuáº©n LaTeX theo máº«u ex_test dÆ°á»›i Ä‘Ã¢y:
YÃªu cáº§u:
- KhÃ´ng thÃªm 'CÃ¢u 1.' hoáº·c 'CÃ¢u 2.'.
- Náº¿u cÃ³ cÃ¡c lá»±a chá»n tráº¯c nghiá»‡m (A., B., C., D.), hÃ£y chuyá»ƒn thÃ nh:
  \\choice
  {{Ä‘Ã¡p Ã¡n 1}}
  {{Ä‘Ã¡p Ã¡n 2}}
  {{Ä‘Ã¡p Ã¡n 3}}
  {{Ä‘Ã¡p Ã¡n 4}}
   (má»—i Ä‘Ã¡p Ã¡n trÃªn 1 dÃ²ng riÃªng)
- Náº¿u lÃ  bÃ i Ä‘Ãºng/sai, dÃ¹ng:
  \\choiceTF
  {{má»‡nh Ä‘á» 1}}
  {{má»‡nh Ä‘á» 2}}
  {{má»‡nh Ä‘á» 3}}
  {{má»‡nh Ä‘á» 4}}
- Má»—i bÃ i Ä‘áº·t trong \\begin{{ex}} ... \\end{{ex}}, cÃ³ \\loigiai{{...}} á»Ÿ cuá»‘i.
"""

            try:
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                    #model="llama-3.1-8b-instant",
                    temperature=0.7,
                )
                output = chat_completion.choices[0].message.content.strip()
                st.code(output, language="latex")

                # âœ… TÃ¡ch cÃ¡c cÃ¢u há»i thÃ nh danh sÃ¡ch
                split_questions = split_ex_blocks(output)
                if not split_questions:
                    split_questions = [output]  # fallback náº¿u khÃ´ng tÃ¡ch Ä‘Æ°á»£c

                st.session_state.all_questions = split_questions
                st.success("âœ… HoÃ n táº¥t xá»­ lÃ½ vÄƒn báº£n.")

            except Exception as e:
                st.error(f"Lá»—i khi gá»i Groq API: {e}")

# =========================
# ğŸš€ Sinh cÃ¢u há»i (2 cháº¿ Ä‘á»™ Ä‘áº§u)
# =========================
col_gen = st.columns([1,1,1])
with col_gen[0]:
    submitted = st.button("ğŸš€ Sinh cÃ¢u há»i")
with col_gen[1]:
    export_word_btn = st.button("â¬‡ï¸ Xuáº¥t Word")
with col_gen[2]:
    export_tex_btn = st.button("â¬‡ï¸ Xuáº¥t LaTeX")

if "all_questions" not in st.session_state:
    st.session_state.all_questions = []

if submitted and not mode.startswith("ğŸ“¤"):
    client = Groq(api_key=api_key)
    all_questions = []
    if mode.startswith("âœï¸"):
        if not user_input.strip():
            st.warning("âš ï¸ Vui lÃ²ng nháº­p Ã­t nháº¥t má»™t cÃ¢u há»i máº«u.")
        else:
            prompt = f"""
DÆ°á»›i Ä‘Ã¢y lÃ  cÃ¢u há»i máº«u theo chuáº©n ex_test:
{user_input}

HÃ£y sinh thÃªm {so_luong_tu_nhap} cÃ¢u há»i tÆ°Æ¡ng tá»± báº±ng tiáº¿ng Viá»‡t.
YÃªu cáº§u:
- Giá»¯ nguyÃªn cáº¥u trÃºc LaTeX (\\begin{{ex}} ... \\end{{ex}})
- Náº¿u cÃ¢u máº«u cÃ³ \\choiceTF thÃ¬ sinh Ä‘Ãºng dáº¡ng Ä‘Ã³, náº¿u cÃ³ \\shortans thÃ¬ sinh tÆ°Æ¡ng á»©ng
- Má»—i cÃ¢u cÃ³ \\loigiai{{...}} á»Ÿ cuá»‘i
âš ï¸ Chá»‰ tráº£ vá» LaTeX, khÃ´ng thÃªm chÃº thÃ­ch nÃ o khÃ¡c.
"""
            try:
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                    #model="llama-3.1-8b-instant",
                    temperature=0.7,
                )
                output = chat_completion.choices[0].message.content.strip()
                all_questions.append(output)
                st.code(output, language="latex")
                st.success(f"âœ… ÄÃ£ sinh {so_luong_tu_nhap} cÃ¢u tá»« ná»™i dung nháº­p thá»§ cÃ´ng.")
            except Exception as e:
                st.error(f"Lá»—i khi gá»i Groq API: {e}")
    else:
        for cfg in st.session_state.configs:
            file_path = get_sample_file(cfg["mon"], cfg["lop"], cfg["topic"], cfg["dang_cauhoi"], cfg["muc_do"], cfg["dang"])
            if not os.path.exists(file_path):
                st.warning(f"âŒ KhÃ´ng tÃ¬m tháº¥y file: {file_path}")
                continue
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            cau_truc = "luÃ´n dÃ¹ng \\choiceTF" if "\\choiceTF" in content else ("luÃ´n dÃ¹ng \\shortans" if "\\shortans" in content else "luÃ´n dÃ¹ng \\choice")
            prompt = f"""
ÄÃ¢y lÃ  cÃ¡c cÃ¢u há»i máº«u theo chuáº©n ex_test:
{content}
HÃ£y sinh {cfg['count']} cÃ¢u há»i tÆ°Æ¡ng tá»± báº±ng tiáº¿ng Viá»‡t.
YÃªu cáº§u:
- DÃ¹ng \\begin{{ex}} ... \\end{{ex}}
- {cau_truc}
- Má»—i cÃ¢u cÃ³ \\loigiai{{...}}
- Náº¿u cÃ³ hÃ¬nh tikz thÃ¬ sinh code tikz phÃ¹ há»£p
âš ï¸ Chá»‰ tráº£ vá» LaTeX, khÃ´ng thÃªm chá»¯ nÃ o khÃ¡c.
"""
            try:
                chat_completion = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt}],
                    model="llama-3.3-70b-versatile",
                    #model="llama-3.1-8b-instant",
                    temperature=0.7,
                )
                output = chat_completion.choices[0].message.content.strip()
                st.code(output, language="latex")
                all_questions.append(output)
                st.success(f"âœ… ÄÃ£ sinh {cfg['count']} cÃ¢u tá»« file.")
            except Exception as e:
                st.error(f"Lá»—i khi gá»i Groq API: {e}")

    st.session_state.all_questions = all_questions

# =========================
# ğŸ’¾ Xuáº¥t file
# =========================
if export_word_btn and st.session_state.all_questions:
    word_file = export_word_ex(st.session_state.all_questions, "de_kiem_tra.docx")
    with open(word_file, "rb") as f:
        st.download_button("â¬‡ï¸ Táº£i Word", f, file_name="de_kiem_tra.docx")

if export_tex_btn and st.session_state.all_questions:
    tex_file = export_latex_ex(st.session_state.all_questions, "de_kiem_tra.tex")
    with open(tex_file, "rb") as f:
        st.download_button("â¬‡ï¸ Táº£i LaTeX", f, file_name="de_kiem_tra.tex")

# =========================
# ğŸ‘€ Preview

st.markdown("## ğŸ² Táº¡o mÃ£ Ä‘á» trá»™n tá»± Ä‘á»™ng")

num_versions = st.number_input("Sá»‘ mÃ£ Ä‘á» muá»‘n táº¡o", 1, 10, 3)
mix_questions = st.button("ğŸ”€ Trá»™n vÃ  táº¡o mÃ£ Ä‘á»")

# =========================
# ğŸ§© HÃ m phá»¥ trá»£
# =========================
def shuffle_choices(q_text):
    """Trá»™n ngáº«u nhiÃªn cÃ¡c lá»±a chá»n \\choice{} trong 1 cÃ¢u há»i"""
    pattern = r"\\choice(.*?)(?=\\loigiai|\\end\{ex\})"
    match = re.search(pattern, q_text, re.S)
    if not match:
        return q_text, None  # khÃ´ng cÃ³ lá»±a chá»n

    block = match.group(1)
    lines = [l.strip() for l in block.splitlines() if l.strip()]
    clean_lines = []
    for l in lines:
        is_true = "\\True" in l
        l = l.replace("\\True", "").strip("{} ")
        clean_lines.append((l, is_true))

    random.shuffle(clean_lines)
    new_block = "\\choice\n" + "\n".join(
        "{" + (("\\True " if is_true else "") + l) + "}" for l, is_true in clean_lines
    )

    q_text_new = q_text.replace(match.group(0), new_block)
    new_answer = chr(65 + [i for i, (_, is_true) in enumerate(clean_lines) if is_true][0])
    return q_text_new, new_answer


def classify_question(q_text):
    """XÃ¡c Ä‘á»‹nh loáº¡i cÃ¢u há»i: 4 lá»±a chá»n, Ä‘Ãºng sai, tráº£ lá»i ngáº¯n"""
    if "\\choiceTF" in q_text:
        return "TF"
    elif "\\choice" in q_text:
        return "MC"
    elif "\\shortans" in q_text:
        return "SA"
    else:
        return "OTHER"


# =========================
# ğŸš€ Trá»™n Ä‘á»
# =========================
if mix_questions and st.session_state.all_questions:
    all_q = "\n".join(st.session_state.all_questions)
    questions = split_ex_blocks(all_q)

    os.makedirs("tmp", exist_ok=True)

    word_files, tex_files = [], []

    for ver in range(1, int(num_versions) + 1):
        q_copy = copy.deepcopy(questions)
        random.shuffle(q_copy)

        # --- PhÃ¢n loáº¡i cÃ¢u há»i ---
        q_mc = [q for q in q_copy if classify_question(q) == "MC"]
        q_tf = [q for q in q_copy if classify_question(q) == "TF"]
        q_sa = [q for q in q_copy if classify_question(q) == "SA"]

        # Trá»™n thá»© tá»± trong tá»«ng nhÃ³m
        random.shuffle(q_mc)
        random.shuffle(q_tf)
        random.shuffle(q_sa)

        mixed_questions = []
        answer_key = []

        # --- Pháº§n I: Tráº¯c nghiá»‡m 4 lá»±a chá»n ---
        mixed_questions.append("\\section*{Pháº§n I â€“ Tráº¯c nghiá»‡m 4 lá»±a chá»n}")
        for i, q in enumerate(q_mc, 1):
            q_new, ans = shuffle_choices(q)
            mixed_questions.append(q_new)
            if ans:
                answer_key.append(f"CÃ¢u {i}: {ans}")
            else:
                answer_key.append(f"CÃ¢u {i}: ---")

        # --- Pháº§n II: Tráº¯c nghiá»‡m Ä‘Ãºng sai ---
        start_tf = len(answer_key) + 1
        mixed_questions.append("\\section*{Pháº§n II â€“ Tráº¯c nghiá»‡m Ä‘Ãºng sai}")
        for j, q in enumerate(q_tf, start=start_tf):
            mixed_questions.append(q)
            match = re.findall(r"\\True|\\False", q)
            if match:
                key = " / ".join(match)
                answer_key.append(f"CÃ¢u {j}: {key}")
            else:
                answer_key.append(f"CÃ¢u {j}: ---")

        # --- Pháº§n III: Tráº£ lá»i ngáº¯n ---
        start_sa = len(answer_key) + 1
        mixed_questions.append("\\section*{Pháº§n III â€“ Tráº£ lá»i ngáº¯n}")
        for k, q in enumerate(q_sa, start=start_sa):
            mixed_questions.append(q)
            sa = re.search(r"\\shortans\{(.*?)\}", q)
            if sa:
                answer_key.append(f"CÃ¢u {k}: {sa.group(1).strip()}")
            else:
                answer_key.append(f"CÃ¢u {k}: ---")

        # --- Xuáº¥t Word vÃ  LaTeX ---
        de_file = f"tmp/De_so_{ver}.docx"
        dap_an_file = f"tmp/Dapan_so_{ver}.docx"
        export_word_ex(mixed_questions, de_file)

        doc_ans = Document()
        doc_ans.add_heading(f"ÄÃP ÃN - MÃƒ Äá»€ {ver}", 0)
        for line in answer_key:
            doc_ans.add_paragraph(line)
        doc_ans.save(dap_an_file)
        word_files += [de_file, dap_an_file]

        tex_file = f"tmp/De_so_{ver}.tex"
        export_latex_ex(mixed_questions, tex_file)
        dap_an_tex = f"tmp/Dapan_so_{ver}.txt"
        with open(dap_an_tex, "w", encoding="utf-8") as f:
            f.write("\n".join(answer_key))
        tex_files += [tex_file, dap_an_tex]

    # --- ÄÃ³ng gÃ³i ZIP ---
    word_zip = "tmp/De_Word.zip"
    tex_zip = "tmp/De_LaTeX.zip"
    with zipfile.ZipFile(word_zip, "w") as zipf:
        for f in word_files:
            zipf.write(f, os.path.basename(f))
    with zipfile.ZipFile(tex_zip, "w") as zipf:
        for f in tex_files:
            zipf.write(f, os.path.basename(f))

    # LÆ°u Ä‘á»ƒ khÃ´ng máº¥t khi rerun
    st.session_state.word_zip = word_zip
    st.session_state.tex_zip = tex_zip

    st.success(f"âœ… ÄÃ£ táº¡o {num_versions} mÃ£ Ä‘á» vÃ  Ä‘Ã¡p Ã¡n thÃ nh cÃ´ng!")


# =========================
# ğŸ’¾ NÃºt táº£i file ZIP
# =========================
if "word_zip" in st.session_state and os.path.exists(st.session_state.word_zip):
    with open(st.session_state.word_zip, "rb") as f:
        st.download_button("â¬‡ï¸ Táº£i táº¥t cáº£ file Word (.zip)", f, file_name="De_Word.zip")

if "tex_zip" in st.session_state and os.path.exists(st.session_state.tex_zip):
    with open(st.session_state.tex_zip, "rb") as f:
        st.download_button("â¬‡ï¸ Táº£i táº¥t cáº£ file LaTeX (.zip)", f, file_name="De_LaTeX.zip")


    st.success(f"âœ… ÄÃ£ táº¡o {num_versions} mÃ£ Ä‘á» vÃ  Ä‘Ã¡p Ã¡n thÃ nh cÃ´ng!")

# =========================
if st.session_state.all_questions:
    st.markdown("### Xem trÆ°á»›c (5 cÃ¢u Ä‘áº§u)")
    for q in st.session_state.all_questions[:5]:
        st.code(q, language="latex")




